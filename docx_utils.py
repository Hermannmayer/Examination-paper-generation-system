"""
考试试卷生成系统 - Word文档处理工具
"""

from docx import Document
from docx.shared import Pt, Cm, Mm, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from tkinter import filedialog
from core import _format_answers

# 纸张尺寸（mm）
PAPER_SIZES = {
    "A4": (210, 297),
    "A3": (297, 420),
    "B5": (176, 250),
    "Letter": (216, 279),
}


def _mm_to_emu(mm):
    return Mm(mm).emu


def export_to_word(exam_datas, config, exam_count=1):
    """导出试卷到Word文件"""
    if not exam_datas:
        return

    # 纸张尺寸
    paper_name = config.get('paper_size', 'A4')
    pw, ph = PAPER_SIZES.get(paper_name, (210, 297))

    # 页眉文字
    header_text = config.get('header_text', '')
    show_header = bool(header_text)

    basename = "试卷"
    initial = f"{basename}_1-{exam_count}.docx" if exam_count > 1 else f"{basename}.docx"
    output_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word文档", "*.docx")],
        initialfile=initial,
    )
    if not output_path:
        return

    for exam_num, exam_data in enumerate(exam_datas, 1):
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Mm(pw)
        section.page_height = Mm(ph)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # 默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(1)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # ── 页眉 ──
        if show_header:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = hp.add_run(header_text)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # ── 试卷标题 ──
        title_text = f"{exam_data['exam_title']} (试卷{exam_num})" if exam_count > 1 else exam_data['exam_title']
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # ── 考生信息 ──
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(exam_data['student_name'])
        run.font.name = '宋体'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # ── 考试时间 ──
        if exam_data.get('exam_time'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"考试时间: {exam_data['exam_time']} 分钟")
            run.bold = True
            run.font.size = Pt(10)

        # ── 试卷内容 ──
        content_text = exam_data['exam_content']
        lines = content_text.split('\n')
        for line in lines:
            if not line.strip():
                continue

            # 跳过已在页眉中渲染的信息行
            if line.startswith('试卷标题:') or line.startswith('考生信息:'):
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

            # 题型标题行（以4个空格开头，如 "    判断题（..."）
            if line.startswith('    ') and '（' in line:
                run = p.add_run(line.strip())
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                continue

            # 题目行（以"数字. "开头）
            if line[0].isdigit() and '. ' in line[:6]:
                parts = line.split('. ', 1)
                # 题号加粗
                run = p.add_run(f"{parts[0]}.")
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if len(parts) > 1:
                    run = p.add_run(f" {parts[1]}")
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_page_break()

        # ── 答案页 ──
        if exam_data.get('include_answers') and exam_data.get('all_answers'):
            heading = doc.add_heading("参考答案", level=1)
            for run in heading.runs:
                run.font.size = Pt(14)

            all_answers = sorted(exam_data['all_answers'], key=lambda x: x[0])
            ans_list = [a[1] for a in all_answers]

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("全部答案: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(_format_answers(ans_list))

            for label, qtype in [("判断题", "判断题"), ("单选题", "单选题"), ("多选题", "多选题")]:
                filtered = [a[1] for a in all_answers if a[2] == qtype]
                if filtered:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(f"{label}答案: ")
                    run.bold = True
                    run.font.size = Pt(10)
                    p.add_run(_format_answers(filtered))

        # ── 保存 ──
        if exam_count > 1:
            base, ext = os.path.splitext(output_path)
            file_path = f"{base}-{exam_num}{ext}"
        else:
            file_path = output_path

        doc.save(file_path)
